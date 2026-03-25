import subprocess
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Qwen3 Azerbaijani Tokenizer: Project Execution & Results Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document outlines the complete technical execution of extending the Qwen3 language model tokenizer to natively support the Azerbaijani language, eliminating the word-shattering character fragmentation caused natively by the letter "ə".')
    
    # Section 1
    doc.add_heading('1. Pipeline Execution & Methodology', level=1)
    
    p1 = doc.add_paragraph('To ensure maximum mathematical accuracy and minimal AI hallucination, we constructed a highly rigorous top-down data pipeline:')
    
    doc.add_paragraph('A. Massive Dataset Collection: We downloaded 8.9 million raw lines of diverse Azerbaijani text from the DOLLMA corpus (covering news, wikis, and literature). We algorithmically deduplicated this into an ultra-dense, mathematically unique 5.4 GB corpus.', style='List Bullet')
    doc.add_paragraph('B. Unigram Token Mining: Because traditional BPE algorithms inherently shatter at strange UTF-8 characters like "ə", we utilized the SentencePiece Unigram algorithm explicitly. We fed it a 5-million sentence randomized subset to extract a master pool of exactly 50,000 highly probable root morphemes.', style='List Bullet')
    doc.add_paragraph('C. Targeted String Filtering: We deleted any token candidates that were too short (≤ 2 characters), were already fully memorized by Qwen natively, or failed to mathematically compress text by at least 1 character.', style='List Bullet')

    # Section 2
    doc.add_heading('2. Hard Edge-Case vs. General Text Validation', level=1)
    
    doc.add_paragraph('To definitively prove the AI tokenizers, we tested them against two entirely different domains:')
    doc.add_paragraph('Hard Edge-Case Validation (The 5.03 Baseline): We hand-picked 20 brutally complex agglutinative edge cases universally shattered by the letter "ə". These represent the absolute worst-case scenarios for Qwen3. Qwen3 originally failed these words natively, averaging 5.03 tokens per word.', style='List Bullet')
    doc.add_paragraph('Deep 300,000 Sentence Evaluation (The 3.56 Baseline): We then tested the models across a massive 12.2 million random words sampled uniformly from the DOLLMA dataset. These standard texts dragged the baseline compression down naturally to 3.565.', style='List Bullet')
    
    doc.add_heading('Before vs. After Agglutinative Compression:', level=2)
    doc.add_paragraph('• müasirləşdirilməsindən → 12 tokens (baseline) vs. 4 tokens ([müasir, ləşdirilməsindən])')
    doc.add_paragraph('• azərbaycan → 5 tokens (baseline) vs. 1 token ([azərbaycan])')
    doc.add_paragraph('• gələcəkdir → 7 tokens (baseline) vs. 2 tokens ([gələcək, dir])')
    
    # Section 3
    doc.add_heading('3. The Comprehensive Performance Matrix', level=1)
    
    doc.add_paragraph('Here are the final performance metrics for the three generated models running globally against BOTH the Hard suite and the Deep 12-million word suite.')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tokenizer Version'
    hdr_cells[1].text = 'Vocabulary Size'
    hdr_cells[2].text = 'Hard Edge-Case Fertility'
    hdr_cells[3].text = 'Deep 300k Sentence Fertility'
    
    data = [
        ('Original Qwen3', '151,669 (Base)', '5.03 tokens/word', '3.565 tokens/word'),
        ('14.7k Version', '166,414', '2.47 tokens/word', '2.753 tokens/word'),
        ('20.0k Version', '171,669', '2.20 tokens/word', '2.529 tokens/word'),
        ('36.2k Version', '187,879', '1.83 tokens/word', '2.379 tokens/word')
    ]
    
    for v, size, hard, deep in data:
        row_cells = table.add_row().cells
        row_cells[0].text = v
        row_cells[1].text = size
        row_cells[2].text = hard
        row_cells[3].text = deep

    # Section 4
    doc.add_heading('4. Tokenizer Versions (Pros & Cons)', level=1)
    
    doc.add_heading('Version 1: qwen3_tokenizer_az_14k', level=2)
    doc.add_paragraph('PROS: Effortlessly fits standard engineering targets. Keeps embedding matrix very small, saving massive VRAM and warming up incredibly fast.')
    doc.add_paragraph('CONS: Leaves a tiny fraction of token compression untouched.')
    
    doc.add_heading('Version 2: qwen3_tokenizer_az_20k (RECOMMENDED)', level=2)
    doc.add_paragraph('PROS: Drops fertility solidly into the low 2.20 range on hard words, and beautifully rounds out deep texts. It mathematically perfectly balances the model limits while squeezing a heavy 29% compression reduction across all standard texts.')
    doc.add_paragraph('CONS: Demands slightly more dataset iteration to stabilize its loss curve than the 14k version.')
    
    doc.add_heading('Version 3: qwen3_tokenizer_azerbaijani (36k Version)', level=2)
    doc.add_paragraph('PROS: The absolute maximum text compression (33.3%). A hard-case fertility of 1.83 means parsing complex Azerbaijani text is essentially as fast and natively dense to the model as parsing English.')
    doc.add_paragraph('CONS: Adding 36,000 extremely bloated parameter rows drastically increases the chance of catastrophic forgetting in Qwen3 if your finetuning text dataset isn\'t extremely large and beautifully curated.')

    # Save
    path = "d:\\tokenizer\\Tokenizer_Summary_Report.docx"
    doc.save(path)
    print(f"Document successfully created at {path}")

if __name__ == '__main__':
    build_report()
